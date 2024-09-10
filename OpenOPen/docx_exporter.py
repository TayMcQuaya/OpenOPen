from docx import Document
from docx.shared import Pt, RGBColor
from PyQt5.QtGui import QTextCursor,QTextListFormat

def export_to_docx(text_edit, filename):
    doc = Document()
    cursor = text_edit.textCursor()

    for block_num in range(text_edit.document().blockCount()):
        block = text_edit.document().findBlockByNumber(block_num)
        cursor.setPosition(block.position())
        cursor.movePosition(QTextCursor.EndOfBlock, QTextCursor.KeepAnchor)
        
        block_text = cursor.selectedText()

        # Create a new paragraph for each block
        if block.textList():
            list_format = block.textList().format()
            if list_format.style() == QTextListFormat.ListDisc:
                paragraph = doc.add_paragraph(style='List Bullet')
            elif list_format.style() == QTextListFormat.ListDecimal:
                paragraph = doc.add_paragraph(style='List Number')
            else:
                paragraph = doc.add_paragraph()
        else:
            paragraph = doc.add_paragraph()

        # Iterate over each character to apply formatting
        for char_index in range(cursor.selectionStart(), cursor.selectionEnd()):
            cursor.setPosition(char_index)
            cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor)
            char_format = cursor.charFormat()
            run = paragraph.add_run(cursor.selectedText())

            # Set font properties
            run.font.name = char_format.font().family()
            run.font.size = Pt(char_format.font().pointSize())
            
            # Set bold, italic, underline
            run.bold = char_format.font().bold()
            run.italic = char_format.font().italic()
            run.underline = char_format.font().underline()

            # Set text color
            color = char_format.foreground().color()
            run.font.color.rgb = RGBColor(color.red(), color.green(), color.blue())

        paragraph.add_run('\n')  # Add a new line after each block

    doc.save(filename)
