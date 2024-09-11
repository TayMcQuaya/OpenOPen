import os
import sys
from typing import Optional
from docx_exporter import export_to_docx

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QAction, QFileDialog,
    QToolBar, QComboBox, QFontComboBox, QMessageBox, QMenu, QMenuBar, QToolButton,
    QShortcut, QLabel, QDialog, QTextEdit
)
from PyQt5.QtGui import (
    QIcon, QFont, QColor, QPalette, QTextCursor, QTextDocument,
    QKeySequence, QTextListFormat, QTextCharFormat
)
from PyQt5.QtCore import Qt, QSize, QTimer,QFile
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

from zoomable_text_edit import ZoomableTextEdit
from color_wheel import ColorWheel
from find_dialog import FindDialog
from constants import *

class Editor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.dark_mode = False
        self.update_timer = QTimer(self)
        self.update_timer.setSingleShot(True)
        self.update_timer.timeout.connect(self.updateFontControls)
        self.initUI()

    def get_resource_path(self, relative_path):
        """ Get absolute path to resource, works for dev and for PyInstaller """
        base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
        return os.path.join(base_path, relative_path)

    def initUI(self):
        self.setupCentralWidget()
        self.setWindowProperties()
        self.setupActions()
        self.setupMenus()
        self.setupToolbar()
        self.setupShortcuts()
        self.setLightModePalette()
        self.loadStyleSheet(self.get_resource_path('styles/style.qss'))
        self.show()
        self.textEdit.textEdit.cursorPositionChanged.connect(self.updateFontControls)
        self.textEdit.textEdit.textChanged.connect(self.scheduleUpdate)
        self.textEdit.textEdit.cursorPositionChanged.connect(self.scheduleUpdate)

    def setWindowProperties(self):
        self.setWindowTitle('OpenOPen (v.1.0.0)')
        icon_path = self.get_resource_path('app_icon_multi.ico')
        print(f"Editor icon path: {icon_path}")
        print(f"Editor icon exists: {os.path.exists(icon_path)}")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
            print("Icon set successfully in Editor")
        else:
            print(f"Icon file not found in Editor: {icon_path}")
        self.setGeometry(500, 100, 1750, 1200)
    
    def scheduleUpdate(self):
        # Schedule an update in 100ms
        self.update_timer.start(100)
    
    def updateFontControls(self):
        cursor = self.textEdit.textEdit.textCursor()
        char_format = cursor.charFormat()
        block_format = cursor.blockFormat()

        # Update bold button
        is_bold = char_format.fontWeight() == QFont.Bold
        self.boldAction.setChecked(is_bold)
        self.updateButtonStyle(self.boldAction, is_bold)

        # Update italic button
        is_italic = char_format.fontItalic()
        self.italicAction.setChecked(is_italic)
        self.updateButtonStyle(self.italicAction, is_italic)

        # Update underline button
        is_underline = char_format.fontUnderline()
        self.underlineAction.setChecked(is_underline)
        self.updateButtonStyle(self.underlineAction, is_underline)

        # Update alignment buttons
        alignment = block_format.alignment()
        self.alignLeftAction.setChecked(alignment == Qt.AlignLeft)
        self.updateButtonStyle(self.alignLeftAction, alignment == Qt.AlignLeft)
        self.alignCenterAction.setChecked(alignment == Qt.AlignCenter)
        self.updateButtonStyle(self.alignCenterAction, alignment == Qt.AlignCenter)
        self.alignRightAction.setChecked(alignment == Qt.AlignRight)
        self.updateButtonStyle(self.alignRightAction, alignment == Qt.AlignRight)

        # Update list buttons
        current_list = cursor.currentList()
        if current_list:
            list_format = current_list.format()
            is_bullet = list_format.style() == QTextListFormat.ListDisc
            is_numbered = list_format.style() == QTextListFormat.ListDecimal
        else:
            is_bullet = is_numbered = False
        
        self.bulletListAction.setChecked(is_bullet)
        self.updateButtonStyle(self.bulletListAction, is_bullet)
        self.numberedListAction.setChecked(is_numbered)
        self.updateButtonStyle(self.numberedListAction, is_numbered)

        # Update font family and size
        current_font = char_format.font().family()
        current_font_size = char_format.font().pointSize()

        if current_font:
            self.fontFamily.setCurrentFont(QFont(current_font))
        if current_font_size > 0:
            self.fontSize.setCurrentText(str(current_font_size))

    def updateButtonStyle(self, action, is_active):
        for widget in self.toolbar.children():
            if isinstance(widget, QToolButton) and widget.defaultAction() == action:
                if is_active:
                    if self.dark_mode:
                        color = QColor(80, 80, 80)  # Dark gray for dark mode
                    else:
                        color = QColor(224, 224, 224)  # #e0e0e0 for light mode
                    widget.setStyleSheet(f"background-color: {color.name()};")
                else:
                    widget.setStyleSheet("")
                break
    
    def updateAllButtonStyles(self):
        actions_to_update = [
            (self.boldAction, self.textEdit.textEdit.fontWeight() == QFont.Bold),
            (self.italicAction, self.textEdit.textEdit.fontItalic()),
            (self.underlineAction, self.textEdit.textEdit.fontUnderline()),
            (self.alignLeftAction, self.textEdit.textEdit.alignment() == Qt.AlignLeft),
            (self.alignCenterAction, self.textEdit.textEdit.alignment() == Qt.AlignCenter),
            (self.alignRightAction, self.textEdit.textEdit.alignment() == Qt.AlignRight),
        ]

        cursor = self.textEdit.textEdit.textCursor()
        if cursor.currentList():
            list_format = cursor.currentList().format()
            actions_to_update.extend([
                (self.bulletListAction, list_format.style() == QTextListFormat.ListDisc),
                (self.numberedListAction, list_format.style() == QTextListFormat.ListDecimal),
            ])
        else:
            actions_to_update.extend([
                (self.bulletListAction, False),
                (self.numberedListAction, False),
            ])

        for action, is_active in actions_to_update:
            self.updateButtonStyle(action, is_active)
                

    def setupCentralWidget(self):
        centralWidget = QWidget(self)
        self.setCentralWidget(centralWidget)
        layout = QVBoxLayout(centralWidget)
        layout.setContentsMargins(0, 0, 0, 0)
        self.textEdit = ZoomableTextEdit(centralWidget)
        layout.addWidget(self.textEdit)

    def setWindowProperties(self):
        self.setWindowTitle(WINDOW_TITLE)
        icon_path = self.get_resource_path(ICON_PATH)
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        #QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
        #QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps)

        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()
        screen_dpi = screen.logicalDotsPerInch() / 96.0

        window_width = int(screen_geometry.width() * 0.9 * (1.0 / screen_dpi))
        window_height = int(screen_geometry.height() * 0.8 * (1.0 / screen_dpi))

        self.setGeometry(
            (screen_geometry.width() - window_width) // 2,
            (screen_geometry.height() - window_height) // 2,
            window_width,
            window_height
        )

    def setupActions(self):
        self.openAction = QAction(QIcon(self.get_resource_path('icons/open.png')), 'Open', self)
        self.openAction.setShortcut('Ctrl+O')
        self.openAction.triggered.connect(self.openFile)

        self.saveAction = QAction(QIcon(self.get_resource_path('icons/save.png')), 'Save', self)
        self.saveAction.setShortcut('Ctrl+S')
        self.saveAction.triggered.connect(self.saveFile)

        self.docxAction = QAction(QIcon(self.get_resource_path('icons/docx.png')), 'Save as DOCX', self)
        self.docxAction.triggered.connect(self.exportDOCX)

        self.pdfAction = QAction(QIcon(self.get_resource_path('icons/pdf.png')), 'Save as PDF', self)
        self.pdfAction.triggered.connect(self.exportPDF)

        self.printAction = QAction(QIcon(self.get_resource_path('icons/print.png')), 'Print', self)
        self.printAction.setShortcut('Ctrl+P')
        self.printAction.triggered.connect(self.printDocument)

        self.boldAction = QAction(QIcon(self.get_resource_path('icons/bold.png')), 'Bold', self)
        self.boldAction.setShortcut('Ctrl+B')
        self.boldAction.triggered.connect(self.setBold)
        self.boldAction.setCheckable(True)

        self.italicAction = QAction(QIcon(self.get_resource_path('icons/italic.png')), 'Italic', self)
        self.italicAction.setShortcut('Ctrl+I')
        self.italicAction.triggered.connect(self.setItalic)
        self.italicAction.setCheckable(True)

        self.underlineAction = QAction(QIcon(self.get_resource_path('icons/underline.png')), 'Underline', self)
        self.underlineAction.setShortcut('Ctrl+U')
        self.underlineAction.triggered.connect(self.setUnderline)
        self.underlineAction.setCheckable(True)

        self.bulletListAction = QAction(QIcon(self.get_resource_path('icons/bullet_list.png')), 'Bullet List', self)
        self.bulletListAction.triggered.connect(self.toggleBulletList)
        self.bulletListAction.setCheckable(True)

        self.numberedListAction = QAction(QIcon(self.get_resource_path('icons/numbered_list.png')), 'Numbered List', self)
        self.numberedListAction.triggered.connect(self.toggleNumberedList)
        self.numberedListAction.setCheckable(True)

        self.findAction = QAction(QIcon(self.get_resource_path('icons/find.png')), 'Find', self)
        self.findAction.setShortcut('Ctrl+F')
        self.findAction.triggered.connect(self.openFindDialog)

        self.colorAction = QAction(QIcon(self.get_resource_path('icons/color.png')), 'Color', self)
        self.colorAction.triggered.connect(self.setColor)

        self.darkModeAction = QAction(QIcon(self.get_resource_path('icons/dark_mode.png')), 'Toggle Dark Mode', self)
        self.darkModeAction.triggered.connect(self.toggleDarkMode)

        self.undoAction = QAction(QIcon(self.get_resource_path('icons/undo.png')), 'Undo', self)
        self.undoAction.setShortcut('Ctrl+Z')
        self.undoAction.triggered.connect(self.textEdit.undo)

        self.redoAction = QAction(QIcon(self.get_resource_path('icons/redo.png')), 'Redo', self)
        self.redoAction.setShortcut('Ctrl+Y')
        self.redoAction.triggered.connect(self.textEdit.redo)

        self.alignLeftAction = QAction(QIcon(self.get_resource_path('icons/align_left.png')), 'Align Left', self)
        self.alignLeftAction.setShortcut('Ctrl+L')
        self.alignLeftAction.triggered.connect(lambda: self.textEdit.setAlignment(Qt.AlignLeft))
        self.alignLeftAction.setCheckable(True)

        self.alignCenterAction = QAction(QIcon(self.get_resource_path('icons/align_center.png')), 'Align Center', self)
        self.alignCenterAction.setShortcut('Ctrl+E')
        self.alignCenterAction.triggered.connect(lambda: self.textEdit.setAlignment(Qt.AlignCenter))
        self.alignCenterAction.setCheckable(True)

        self.alignRightAction = QAction(QIcon(self.get_resource_path('icons/align_right.png')), 'Align Right', self)
        self.alignRightAction.setShortcut('Ctrl+R')
        self.alignRightAction.triggered.connect(lambda: self.textEdit.setAlignment(Qt.AlignRight))
        self.alignRightAction.setCheckable(True)

        self.zoomInAction = QAction(QIcon(self.get_resource_path('icons/zoom_in.png')), 'Zoom In', self)
        self.zoomInAction.setShortcut('Ctrl++')
        self.zoomInAction.triggered.connect(self.textEdit.zoomIn)

        self.zoomOutAction = QAction(QIcon(self.get_resource_path('icons/zoom_out.png')), 'Zoom Out', self)
        self.zoomOutAction.setShortcut('Ctrl+-')
        self.zoomOutAction.triggered.connect(self.textEdit.zoomOut)

    def setupMenus(self):
        menubar = self.menuBar()
        menubar.setBaseSize(QSize(60, 60))
        
        font = QFont()
        font.setPointSize(14)  
        self.menuBar().setFont(font)

        menubar.setContextMenuPolicy(Qt.CustomContextMenu)
        menubar.customContextMenuRequested.connect(self.showMenuBarContextMenu)

        fileMenu = menubar.addMenu('&File')
        fileMenu.addActions([self.openAction, self.saveAction, self.docxAction, self.pdfAction, self.printAction])

        editMenu = menubar.addMenu('&Edit')
        editMenu.addActions([
            self.boldAction, self.italicAction, self.underlineAction,                              # Text formatting
            self.alignLeftAction, self.alignCenterAction, self.alignRightAction,                   # Alignment actions
            self.bulletListAction, self.numberedListAction,                                        # List formatting
            self.colorAction,                                                                      # Text color
            self.undoAction, self.redoAction,                                                      # Undo/Redo
            self.zoomInAction, self.zoomOutAction,                                                 # Zoom controls
            self.darkModeAction                                                                    # Dark mode toggle
        ])

        viewMenu = self.menuBar().addMenu('&View')
        viewMenu.addActions([self.zoomInAction, self.zoomOutAction])
        
        helpMenu = menubar.addMenu('&Help')
        readmeAction = QAction(QIcon('icons/help.png'), 'Open README', self)
        readmeAction.triggered.connect(self.openReadme)
        helpMenu.addAction(readmeAction)

    def setupToolbar(self):
        self.toolbar = QToolBar()
        self.addToolBar(Qt.TopToolBarArea, self.toolbar)
        self.toolbar.setMovable(False)
        self.toolbar.setIconSize(QSize(50, 50))
        self.toolbar.setStyleSheet("""
            QToolBar {
                spacing: 5px;
                padding: 3px;
            }
            QToolButton {
                min-width: 30px;
                min-height: 30px;
            }
        """)

        actions = [
            self.openAction, self.saveAction, self.docxAction, self.pdfAction, self.printAction,
            self.boldAction, self.italicAction, self.underlineAction,
            self.alignLeftAction, self.alignCenterAction, self.alignRightAction,
            self.bulletListAction, self.numberedListAction,
            self.colorAction,
            self.undoAction, self.redoAction, self.findAction,
            self.zoomInAction, self.zoomOutAction,
        ]

        for action in actions:
            self.toolbar.addAction(action)
            widget = self.toolbar.widgetForAction(action)
            widget.setFixedSize(60, 60)

        # Add dark mode button at the end
        self.toolbar.addAction(self.darkModeAction)
        widget = self.toolbar.widgetForAction(self.darkModeAction)
        widget.setFixedSize(60, 60)

        self.setupFontControls()
        self.setupZoomLabel()

        self.toolbar.setContextMenuPolicy(Qt.CustomContextMenu)
        self.toolbar.customContextMenuRequested.connect(self.showToolbarContextMenu)
        
        self.toolbar.update()
        self.toolbar.repaint()

    def setupFontControls(self):
        self.fontFamily = QFontComboBox(self)
        self.fontFamily.setCurrentFont(QFont("Times New Roman"))
        self.fontFamily.setStyleSheet("font-size: 14pt;")
        self.fontFamily.setFixedSize(200, 40)
        self.fontFamily.currentFontChanged.connect(self.setFontFamily)
        self.toolbar.addWidget(self.fontFamily)

        self.fontSize = QComboBox(self)
        self.fontSize.addItems([str(i) for i in range(8, 73, 2)])
        self.fontSize.setStyleSheet("font-size: 14pt;")
        self.fontSize.setCurrentText('16')
        self.fontSize.setFixedSize(70, 40)
        self.fontSize.currentTextChanged.connect(self.setFontSize)
        self.toolbar.addWidget(self.fontSize)

    def setupZoomLabel(self):
        self.toolbar.addAction(self.zoomOutAction)
        self.zoomLabel = QLabel("100%")
        self.zoomLabel.setStyleSheet("font-weight: bold;font-size: 14pt;")
        self.toolbar.addWidget(self.zoomLabel)
        self.toolbar.addAction(self.zoomInAction)

    def setupShortcuts(self):
        QShortcut(QKeySequence('Ctrl+Q'), self, self.close)
        QShortcut(QKeySequence('Ctrl+='), self, self.increaseFontSize)
        QShortcut(QKeySequence('Ctrl+-'), self, self.decreaseFontSize)
        self.textEdit.zoomChanged.connect(self.updateZoomLabel)

    def setLightModePalette(self):
        palette = self.palette()
        palette.setColor(QPalette.Window, Qt.white)
        palette.setColor(QPalette.WindowText, Qt.black)
        palette.setColor(QPalette.Base, Qt.white)
        palette.setColor(QPalette.AlternateBase, Qt.lightGray)
        palette.setColor(QPalette.ToolTipBase, Qt.white)
        palette.setColor(QPalette.ToolTipText, Qt.black)
        palette.setColor(QPalette.Text, Qt.black)
        palette.setColor(QPalette.Button, Qt.lightGray)
        palette.setColor(QPalette.ButtonText, Qt.black)
        palette.setColor(QPalette.BrightText, Qt.red)
        palette.setColor(QPalette.Link, QColor(0, 0, 255))
        palette.setColor(QPalette.Highlight, QColor(0, 120, 215))
        palette.setColor(QPalette.HighlightedText, Qt.white)

        self.setPalette(palette)
        QApplication.instance().setPalette(palette)
        self.textEdit.setPalette(palette)
        self.textEdit.textEdit.setPalette(palette)
        self.textEdit.setBackgroundBrush(palette.color(QPalette.Window))

    def loadStyleSheet(self, path):
        with open(path, 'r') as file:
            stylesheet = file.read()

        custom_styles = """
        QMenuBar {
            padding: 5px;
            margin: 0px;
        }
        QToolBar {
            padding: 5px;
            margin: 0px;
        }
        """
        
        stylesheet += custom_styles
        self.setStyleSheet(stylesheet)

    def showToolbarContextMenu(self, pos):
        contextMenu = QMenu(self)
        toggleAction = contextMenu.addAction("Hide Toolbar" if self.toolbar.isVisible() else "Show Toolbar")
        action = contextMenu.exec_(self.toolbar.mapToGlobal(pos))
        if action == toggleAction:
            self.toolbar.setVisible(not self.toolbar.isVisible())

    def showMenuBarContextMenu(self, pos):
        contextMenu = QMenu(self)
        toggleAction = contextMenu.addAction("Hide Toolbar" if self.toolbar.isVisible() else "Show Toolbar")
        action = contextMenu.exec_(self.menuBar().mapToGlobal(pos))
        if action == toggleAction:
            self.toolbar.setVisible(not self.toolbar.isVisible())

    def updateZoomLabel(self, factor):
        self.zoomLabel.setText(f"{factor}%")

    def increaseFontSize(self):
        currentSize = self.textEdit.textEdit.fontPointSize()
        self.textEdit.textEdit.setFontPointSize(currentSize + 1)

    def decreaseFontSize(self):
        currentSize = self.textEdit.textEdit.fontPointSize()
        if currentSize > 1:
            self.textEdit.textEdit.setFontPointSize(currentSize - 1)

    def toggleDarkMode(self):
        self.dark_mode = not self.dark_mode
        self.updateIcons(dark=self.dark_mode)

        style_path = self.get_resource_path('styles/dark_style.qss' if self.dark_mode else 'styles/style.qss')
        self.loadStyleSheet(style_path)
        
        
        palette = self.palette()
        if self.dark_mode:
            palette.setColor(QPalette.Window, QColor(53, 53, 53))
            palette.setColor(QPalette.WindowText, Qt.white)
            palette.setColor(QPalette.Base, QColor(53, 53, 53))
            palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
            palette.setColor(QPalette.ToolTipBase, Qt.white)
            palette.setColor(QPalette.ToolTipText, Qt.white)
            palette.setColor(QPalette.Text, Qt.white)
            palette.setColor(QPalette.Button, QColor(53, 53, 53))
            palette.setColor(QPalette.ButtonText, Qt.white)
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Link, QColor(42, 130, 218))
            palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
            palette.setColor(QPalette.HighlightedText, Qt.black)
        else:
            palette.setColor(QPalette.Window, Qt.white)
            palette.setColor(QPalette.WindowText, Qt.black)
            palette.setColor(QPalette.Base, Qt.white)
            palette.setColor(QPalette.AlternateBase, Qt.lightGray)
            palette.setColor(QPalette.ToolTipBase, Qt.white)
            palette.setColor(QPalette.ToolTipText, Qt.black)
            palette.setColor(QPalette.Text, Qt.black)
            palette.setColor(QPalette.Button, Qt.lightGray)
            palette.setColor(QPalette.ButtonText, Qt.black)
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Link, QColor(0, 0, 255))
            palette.setColor(QPalette.Highlight, QColor(0, 120, 215))
            palette.setColor(QPalette.HighlightedText, Qt.white)
        
        self.setPalette(palette)
        QApplication.instance().setPalette(palette)
        self.textEdit.setPalette(palette)
        self.textEdit.textEdit.setPalette(palette)
        self.textEdit.setBackgroundBrush(palette.color(QPalette.Window))

        self.repaint()
        for widget in self.findChildren(QWidget):
            widget.update()

        self.updateAllButtonStyles()

        menubar = self.menuBar()
        menubar_palette = menubar.palette()
        menubar_palette.setColor(QPalette.ButtonText, Qt.white if self.dark_mode else Qt.black)
        menubar.setPalette(menubar_palette)
        
        self.zoomLabel.setStyleSheet(f"font-weight: bold; color: {'white' if self.dark_mode else 'black'};font-size: 14pt;")
        self.fontFamily.setStyleSheet(f"color: {'white' if self.dark_mode else 'black'};")
        self.fontSize.setStyleSheet(f"color: {'white' if self.dark_mode else 'black'};")

    def updateIcons(self, dark):
        icon_suffix = '_dark' if dark else ''
        icons_to_update = {
            self.printAction: 'print',
            self.boldAction: 'bold',
            self.italicAction: 'italic',
            self.underlineAction: 'underline',
            self.darkModeAction: 'dark_mode',
            self.alignLeftAction: 'align_left',
            self.alignRightAction: 'align_right',
            self.alignCenterAction: 'align_center',
            self.bulletListAction: 'bullet_list',
            self.numberedListAction: 'numbered_list'
        }
        for action, icon_name in icons_to_update.items():
            action.setIcon(QIcon(self.get_resource_path(f'icons/{icon_name}{icon_suffix}.png')))

    def openReadme(self):
        readme_path = self.get_resource_path(os.path.join('resources', 'README.txt'))
        if os.path.exists(readme_path):
            try:
                # Try UTF-8 encoding first
                with open(readme_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            except UnicodeDecodeError:
                try:
                    # If UTF-8 fails, try UTF-16
                    with open(readme_path, 'r', encoding='utf-16') as file:
                        content = file.read()
                except UnicodeDecodeError:
                    # If both fail, use latin-1 which should read all bytes
                    with open(readme_path, 'r', encoding='latin-1') as file:
                        content = file.read()
            
            dlg = QDialog(self)
            dlg.setWindowTitle("README")
            layout = QVBoxLayout()
            textEdit = QTextEdit()
            textEdit.setText(content)
            textEdit.setReadOnly(True)
            layout.addWidget(textEdit)
            dlg.setLayout(layout)
            dlg.resize(self.size())
            dlg.exec_()
        else:
            QMessageBox.warning(self, "Error", "README file not found.")
    def openFile(self):
        fname, _ = QFileDialog.getOpenFileName(self, 'Open file', '/', "Rich Text Files (*.rtf);;Word Documents (*.docx);;All Files (*)")
        if fname:
            if fname.lower().endswith('.docx'):
                self.openDocx(fname)
            elif fname.lower().endswith('.rtf'):
                self.openRtf(fname)
            else:
                with open(fname, 'r') as file:
                    self.textEdit.setText(file.read())

    def openDocx(self, fname):
        try:
            doc = Document(fname)
            self.textEdit.textEdit.clear()
            cursor = self.textEdit.textEdit.textCursor()
            
            for para in doc.paragraphs:
                if para.style.name.startswith('List'):
                    list_format = QTextListFormat()
                    if 'Bullet' in para.style.name:
                        list_format.setStyle(QTextListFormat.ListDisc)
                    elif 'Number' in para.style.name:
                        list_format.setStyle(QTextListFormat.ListDecimal)
                    cursor.insertList(list_format)
                else:
                    if cursor.position() != 0:
                        cursor.insertBlock()

                for run in para.runs:
                    char_format = QTextCharFormat()
                    
                    # Font family
                    if run.font.name:
                        char_format.setFontFamily(run.font.name)
                    
                    # Font size
                    if run.font.size:
                        char_format.setFontPointSize(run.font.size.pt)
                    
                    # Bold
                    char_format.setFontWeight(QFont.Bold if run.bold else QFont.Normal)
                    
                    # Italic
                    italic_value = False if run.italic is None else run.italic
                    char_format.setFontItalic(italic_value)
                    
                    # Underline
                    underline_value = False if run.underline is None else run.underline
                    char_format.setFontUnderline(underline_value)
                    
                    # Color
                    if run.font.color and run.font.color.rgb:
                        r, g, b = run.font.color.rgb
                        color = QColor(r, g, b)
                        char_format.setForeground(color)
                    
                    cursor.setCharFormat(char_format)
                    cursor.insertText(run.text)

            self.textEdit.textEdit.setTextCursor(cursor)
            self.textEdit.textEdit.document().setModified(False)
            
        except Exception as e:
            QMessageBox.warning(self, "Open Error", f"Failed to open DOCX file: {str(e)}")
            print(f"Error details: {str(e)}")  # For debugging


    def openRtf(self, fname):
        try:
            with open(fname, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_text = file.read()
            document = QTextDocument()
            document.setHtml(rtf_text)
            self.textEdit.setDocument(document)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to open RTF file: {str(e)}")

    def saveFile(self):
        fname, _ = QFileDialog.getSaveFileName(self, 'Save file', '/', "Rich Text Files (*.rtf)")
        if fname:
            if not fname.lower().endswith('.rtf'):
                fname += '.rtf'
            try:
                with open(fname, 'w', encoding='utf-8') as file:
                    file.write(self.textEdit.toHtml())
                QMessageBox.information(self, "Save Successful", "File saved successfully.")
            except Exception as e:
                QMessageBox.warning(self, "Save Error", f"Failed to save the file: {str(e)}")

    def exportDOCX(self):
        fname, _ = QFileDialog.getSaveFileName(self, 'Export DOCX', '/', filter="Word Documents (*.docx)")
        if fname:
            try:
                # Call the export_to_docx function from docx_exporter.py
                export_to_docx(self.textEdit.textEdit, fname)
                QMessageBox.information(self, "Export Successful", "File exported successfully.")
            except Exception as e:
                QMessageBox.warning(self, "Export Error", f"Failed to export DOCX: {str(e)}")

    def exportPDF(self):
        fname, _ = QFileDialog.getSaveFileName(self, 'Export PDF', '/', filter="PDF Files (*.pdf)")
        if fname:
            try:
                printer = QPrinter(QPrinter.HighResolution)
                printer.setOutputFormat(QPrinter.PdfFormat)
                printer.setOutputFileName(fname)
                self.textEdit.textEdit.document().print_(printer)
                QMessageBox.information(self, "Export Successful", "File exported successfully.")
            except Exception as e:
                QMessageBox.warning(self, "Export Error", f"Failed to export PDF: {str(e)}")

    def printDocument(self):
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            self.textEdit.textEdit.print_(printer)

    def setBold(self):
        self.textEdit.textEdit.setFontWeight(QFont.Bold if self.textEdit.textEdit.fontWeight() != QFont.Bold else QFont.Normal)

    def setItalic(self):
        self.textEdit.textEdit.setFontItalic(not self.textEdit.textEdit.fontItalic())

    def setUnderline(self):
        self.textEdit.textEdit.setFontUnderline(not self.textEdit.textEdit.fontUnderline())

    def setColor(self):
        color_wheel = ColorWheel(self)
        if color_wheel.exec_():
            color = color_wheel.selectedColor()
            if color.isValid():
                self.textEdit.textEdit.setTextColor(color)

    def setFontFamily(self, font):
        cursor = self.textEdit.textEdit.textCursor()

        # If there's a selection, apply the font to the selected text
        if cursor.hasSelection():
            char_format = QTextCharFormat()
            char_format.setFontFamily(font.family())
            cursor.mergeCharFormat(char_format)
        else:
            # If no selection, apply the font to future text
            self.textEdit.textEdit.setFontFamily(font.family())
    def setFontSize(self, size):
        cursor = self.textEdit.textEdit.textCursor()
        
        # If there's a selection, only apply the size to the selection
        if cursor.hasSelection():
            char_format = QTextCharFormat()
            char_format.setFontPointSize(float(size))
            cursor.mergeCharFormat(char_format)
        else:
            # If no selection, only set the font size for future text
            self.textEdit.textEdit.setFontPointSize(float(size))


    def toggleBulletList(self):
        cursor = self.textEdit.textEdit.textCursor()
        if cursor.currentList():
            list_item = cursor.currentList()
            list_item.remove(cursor.block())
            block_format = cursor.blockFormat()
            block_format.setIndent(0)
            cursor.setBlockFormat(block_format)
        else:
            list_format = QTextListFormat()
            list_format.setStyle(QTextListFormat.ListDisc)
            cursor.createList(list_format)
        self.textEdit.textEdit.setTextCursor(cursor)

    def toggleNumberedList(self):
        cursor = self.textEdit.textEdit.textCursor()
        if cursor.currentList():
            list_item = cursor.currentList()
            list_item.remove(cursor.block())
            block_format = cursor.blockFormat()
            block_format.setIndent(0)
            cursor.setBlockFormat(block_format)
        else:
            list_format = QTextListFormat()
            list_format.setStyle(QTextListFormat.ListDecimal)
            cursor.createList(list_format)
        self.textEdit.textEdit.setTextCursor(cursor)

    def openFindDialog(self):
        self.findDialog = FindDialog(self)
        self.findDialog.findBtn.clicked.connect(self.findNext)
        self.findDialog.show()

    def findNext(self):
        searchText = self.findDialog.input.text()
        if not searchText:
            return
        cursor = self.textEdit.textEdit.textCursor()
        document = self.textEdit.textEdit.document()
        cursor = document.find(searchText, cursor)
        if cursor.isNull():
            cursor = document.find(searchText, QTextCursor(document.begin()))
        if not cursor.isNull():
            self.textEdit.textEdit.setTextCursor(cursor)
        else:
            self.statusBar().showMessage(f"Cannot find '{searchText}'", 2000)

